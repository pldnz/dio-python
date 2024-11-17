import requests
from docx import Document
import os

subscription_key = "YOUR_SUBSCRIPTION_KEY"
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = "eastus"

target_language = "pt-br"

def translator_text(text, target_language=target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
    'Ocp-Apim-Subscription-Key': subscription_key,
    'Ocp-Apim-Subscription-Region': location,
    'Content-type': 'application/json',
    'X-ClientTraceId': str(os.urandom(16))
  }
  body = [{
    'text': text
  }]
  params = {
    'api-version': '3.0',
    'from': 'en',
    'to': target_language
  }
  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  # print(response)
  return response[0]["translations"][0]["text"]


def translate_document(path):
  document = Document(path)
  full_text = []
  for para in document.paragraphs:
    translated_text = translator_text(para.text)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)
  path_translated = path.replace(".docx", f"_{target_language}.docx")
  translated_doc.save(path_translated)

  return path_translated


translator_text("Sheets of empty canvas Untouched sheets of clay Were laid spread out before me As her body once did")
